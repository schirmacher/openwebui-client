#!/usr/bin/env python3
"""
Tutorial for using the OpenWebUIClient library, demonstrating a wide range of use cases.
"""
import os
import re
import sys
import json
from PIL import Image, ImageDraw
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
# Import the client library
from openwebui_client import OpenWebUIClient, OpenWebUIMessageBuilder

# --- Configuration ---
# Load settings from environment variables
BASE_URL = os.getenv("OPENWEBUI_URL")
API_KEY = os.getenv("OPENWEBUI_API_KEY")
MODEL = os.getenv("OPENWEBUI_MODEL")
MULTIMODAL_MODEL = os.getenv("OPENWEBUI_MULTIMODAL_MODEL")


# --- Helper Functions ---
def print_header(title):
    """Prints a formatted header for each use case."""
    print("\n" + "=" * 80)
    print(f"| {title:^76} |")
    print("=" * 80 + "\n")


def create_dummy_image(file_path: str, shape: str, color: str):
    """Creates a simple image with a geometric shape."""
    img = Image.new('RGB', (200, 200), color='white')
    draw = ImageDraw.Draw(img)
    if shape == "square":
        draw.rectangle([50, 50, 150, 150], fill=color)
    elif shape == "circle":
        draw.ellipse([50, 50, 150, 150], fill=color)
    elif shape == "triangle":
        draw.polygon([(50, 150), (100, 50), (150, 150)], fill=color)
    img.save(file_path)
    print(f"Generated dummy image: {file_path}")


def create_dummy_pdf(file_path: str, title: str, text_lines: list):
    """Creates a simple PDF document."""
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    text = c.beginText(72, height - 72)
    text.setFont("Helvetica-Bold", 14)
    text.textLine(title)
    text.setFont("Helvetica", 12)
    text.moveCursor(0, 14)  # Add a space
    for line in text_lines:
        text.textLine(line)
    c.drawText(text)
    c.save()
    print(f"Generated dummy PDF: {file_path}")


def create_dummy_docx(file_path: str, title: str, text_lines: list):
    """Creates a simple DOCX document."""
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text_lines:
        doc.add_paragraph(line)
    doc.save(file_path)
    print(f"Generated dummy DOCX: {file_path}")


def stream_and_print_response(stream_generator, wrap_width=120):
    """
    Streams a response, prints it with live line-wrapping, and returns the full content.
    """
    full_response = ""
    current_column = 0
    for chunk in stream_generator:
        # Gracefully handle chunks that don't contain 'choices'.
        # These are often metadata messages from the server when using RAG (file uploads).
        if 'choices' not in chunk or not chunk['choices']:
            continue

        delta = chunk['choices'][0].get('delta', {})
        content = delta.get('content')

        if content:
            full_response += content
            for char in content:
                print(char, end="", flush=True)
                if char == '\n':
                    current_column = 0
                else:
                    current_column += 1
                if current_column >= wrap_width:
                    print()
                    current_column = 0
    print()
    return full_response


def print_source_references_if_any(response_text: str, uploaded_files: list):
    """
    Checks for [n] style citations in the response and prints the corresponding source files.
    """
    citations = re.findall(r'\[\d+\]', response_text)
    if not citations:
        return

    print("\n--- Source Documents ---")
    source_map = {f"[{i + 1}]": f['filename'] for i, f in enumerate(uploaded_files)}
    unique_citations_in_text = sorted(list(set(citations)), key=lambda x: int(x.strip('[]')))

    for citation in unique_citations_in_text:
        if citation in source_map:
            print(f"{citation}: {source_map[citation]}")
    print("------------------------")


# --- Use Cases ---

def use_case_1_list_models(client: OpenWebUIClient):
    """1. Query open-webui for the list of existing models."""
    print_header("Use Case 1: List Available Models")
    try:
        models = client.get_models()
        if not models:
            print("No models found.")
            return
        print("Available models:")
        for model in models:
            print(f"- {model.get('id', 'N/A')}")
    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_2_simple_query(client: OpenWebUIClient):
    """2. Make a simple query and write out the result. (Non-streaming)"""
    print_header("Use Case 2: Simple Chat Completion (Non-Streaming)")
    try:
        prompt = "What are the three main benefits of using Python?"
        print(f"User Query: {prompt}\n")

        builder = OpenWebUIMessageBuilder()
        builder.add_user_message(prompt)
        messages = builder.build()

        response = client.chat_completion(
            model=MODEL,
            messages=messages
        )

        assistant_response = response['choices'][0]['message']['content']
        print("Assistant Response:\n---")
        print(assistant_response)
        print("---")
    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_3_streaming_query(client: OpenWebUIClient):
    """3. Make a query in streaming mode with live line-wrapping."""
    print_header("Use Case 3: Streaming Chat Completion")
    try:
        prompt = "Explain the concept of 'dark matter' to a 10-year-old in a few paragraphs."
        print(f"User Query: {prompt}\n")

        builder = OpenWebUIMessageBuilder()
        builder.add_user_message(prompt)
        messages = builder.build()

        stream = client.stream_chat_completion(
            model=MODEL,
            messages=messages
        )

        print("Assistant Response (streaming):\n---")
        stream_and_print_response(stream)
        print("---")

    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_4_conversational_context(client: OpenWebUIClient):
    """4. Demonstrate a conversation with context using streaming."""
    print_header("Use Case 4: Streaming Conversational Context")
    try:
        builder = OpenWebUIMessageBuilder(
            system_prompt="You are a helpful travel assistant."
        )

        prompt1 = "I want to plan a 3-day trip to Rome. What are three must-see landmarks?"
        print(f"User Query 1: {prompt1}\n")
        builder.add_user_message(prompt1)

        stream1 = client.stream_chat_completion(MODEL, builder.build())
        print("Assistant Response 1 (streaming):\n---")
        answer1 = stream_and_print_response(stream1)
        print("---\n")
        builder.add_assistant_message(answer1)

        prompt2 = "Great. Can you suggest a good place to eat near the first landmark you mentioned?"
        print(f"User Query 2: {prompt2}\n")
        builder.add_user_message(prompt2)

        stream2 = client.stream_chat_completion(MODEL, builder.build())
        print("Assistant Response 2 (streaming):\n---")
        stream_and_print_response(stream2)
        print("---")

    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_5_system_prompt_persona(client: OpenWebUIClient):
    """5. Demonstrate how a system prompt sets the model's persona."""
    print_header("Use Case 5: Setting Model Persona with a System Prompt")
    prompt = "Tell me about the moon."
    print(f"User Query (for both scenarios): {prompt}\n")

    print("--- Scenario 1: Standard Assistant ---")
    builder1 = OpenWebUIMessageBuilder(system_prompt="You are a helpful and factual scientific assistant.")
    builder1.add_user_message(prompt)
    stream1 = client.stream_chat_completion(MODEL, builder1.build())
    stream_and_print_response(stream1)

    print("\n--- Scenario 2: Pirate Persona ---")
    builder2 = OpenWebUIMessageBuilder(
        system_prompt="You are a salty pirate captain who begrudgingly answers questions.")
    builder2.add_user_message(prompt)
    stream2 = client.stream_chat_completion(MODEL, builder2.build())
    stream_and_print_response(stream2)
    print("---")


def use_case_6_robust_json_extraction(client: OpenWebUIClient):
    """6. Demonstrate a robust method for getting JSON output from a model."""
    print_header("Use Case 6: Robust JSON Extraction")
    print("This method uses strong prompting and client-side extraction to reliably get JSON.\n")
    try:
        # 1. Use a very strict system prompt to force the model's behavior.
        builder = OpenWebUIMessageBuilder(
            system_prompt="You are a data extraction robot. Your ONLY output is valid JSON. Do not include markdown formatting, introductions, or any conversational text. Your entire response must be a single JSON object."
        )
        # 2. Provide few-shot examples to reinforce the desired format.
        builder.add_user_message("The book 'Dune' was written by Frank Herbert in 1965.")
        builder.add_assistant_message('{"title": "Dune", "author": "Frank Herbert", "year": 1965}')
        builder.add_user_message("Arthur C. Clarke wrote '2001: A Space Odyssey' in 1968.")
        builder.add_assistant_message('{"title": "2001: A Space Odyssey", "author": "Arthur C. Clarke", "year": 1968}')

        final_prompt = "Isaac Asimov's 'Foundation' was published in 1951."
        print(f"Final User Query: {final_prompt}\n")
        builder.add_user_message(final_prompt)

        stream = client.stream_chat_completion(MODEL, builder.build())
        print("Assistant's Full Response (streaming):\n---")
        full_response = stream_and_print_response(stream)
        print("---")

        # 3. Proactively clean the output on the client side.
        # Find the first '{' and the last '}' to extract the JSON block.
        json_str = ""
        match = re.search(r'\{.*\}', full_response, re.DOTALL)
        if match:
            json_str = match.group(0)
            print("\nExtracted JSON block for validation:")
            print(json_str)
        else:
            print("\nCould not find a JSON block in the response.")

        # 4. Validate the extracted string.
        try:
            json.loads(json_str)
            print("\nValidation: SUCCESS - The extracted block is valid JSON.")
        except json.JSONDecodeError:
            print("\nValidation: FAILED - The extracted block is not valid JSON.")
    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_7_json_mode(client: OpenWebUIClient):
    """7. Demonstrate forcing a guaranteed JSON output."""
    print_header("Use Case 7: Guaranteed JSON Output with `response_format`")
    print("This is a more robust way to get structured data than few-shot prompting.\n")
    print("Note that NOT ALL LLMs SUPPORT THIS FORMAT.\n")
    try:
        prompt = "Extract the contact information from this text: 'You can reach Jane Doe, the project manager at Innovate Corp, via j.doe@example.com.'"
        print(f"User Query: {prompt}\n")

        stream = client.stream_chat_completion(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )

        print("Assistant's Guaranteed JSON Response (streaming):\n---")
        full_response = stream_and_print_response(stream)
        print("---")
        try:
            json.loads(full_response)
            print("\nValidation: Successfully parsed response as JSON object.")
        except json.JSONDecodeError:
            print("\nValidation: FAILED to parse the response as JSON.")
    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_8_generation_parameters(client: OpenWebUIClient):
    """8. Demonstrate controlling output with temperature."""
    print_header("Use Case 8: Controlling Generation with Parameters")
    prompt = "Write a short, one-sentence tagline for a new brand of coffee."
    print(f"User Query: {prompt}\n")
    try:
        print("--- Scenario 1: Low Temperature (0.1) for a focused response ---")
        stream1 = client.stream_chat_completion(
            model=MODEL, messages=[{"role": "user", "content": prompt}], temperature=0.1
        )
        stream_and_print_response(stream1)

        print("\n--- Scenario 2: High Temperature (1.0) for a creative response ---")
        stream2 = client.stream_chat_completion(
            model=MODEL, messages=[{"role": "user", "content": prompt}], temperature=1.0
        )
        stream_and_print_response(stream2)
        print("---")
    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_9_single_image_query(client: OpenWebUIClient):
    """9. Make a query about a single image using streaming."""
    print_header("Use Case 9: Single Image Query with Streaming")
    if not MULTIMODAL_MODEL:
        print("Skipping: OPENWEBUI_MULTIMODAL_MODEL not set.")
        return
    try:
        image_path = "image1.png"
        create_dummy_image(image_path, "square", "red")

        prompt = "Describe the object in this image. What is its shape and color?"
        print(f"User Query: {prompt}")
        print(f"Image: {image_path}\n")

        builder = OpenWebUIMessageBuilder()
        builder.add_user_message_with_images(
            text=prompt,
            image_paths=[image_path]
        )

        stream = client.stream_chat_completion(MULTIMODAL_MODEL, builder.build())
        print("Assistant Response (streaming):\n---")
        stream_and_print_response(stream)
        print("---")
        os.remove(image_path)

    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_10_multi_image_query(client: OpenWebUIClient):
    """10. Make a conversational query about multiple images using streaming."""
    print_header("Use Case 10: Multi-Image Streaming Conversational Query")
    if not MULTIMODAL_MODEL:
        print("Skipping: OPENWEBUI_MULTIMODAL_MODEL not set.")
        return
    try:
        img1_path, img2_path = "image1.png", "image2.png"
        create_dummy_image(img1_path, "circle", "blue")
        create_dummy_image(img2_path, "triangle", "green")
        builder = OpenWebUIMessageBuilder()

        prompt1 = "What shapes and colors do you see in these images?"
        print(f"User Query 1: {prompt1}")
        print(f"Images: {img1_path}, {img2_path}\n")
        builder.add_user_message_with_images(
            text=prompt1,
            image_paths=[img1_path, img2_path]
        )

        stream1 = client.stream_chat_completion(MULTIMODAL_MODEL, builder.build())
        print("Assistant Response 1 (streaming):\n---")
        answer1 = stream_and_print_response(stream1)
        print("---\n")
        builder.add_assistant_message(answer1)

        prompt2 = "How many distinct objects are there in total across both images?"
        print(f"User Query 2: {prompt2}\n")
        builder.add_user_message(prompt2)

        stream2 = client.stream_chat_completion(MULTIMODAL_MODEL, builder.build())
        print("Assistant Response 2 (streaming):\n---")
        stream_and_print_response(stream2)
        print("---")

        os.remove(img1_path)
        os.remove(img2_path)

    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_11_single_file_query(client: OpenWebUIClient):
    """11. Upload a PDF and ask a question about its content with citations."""
    print_header("Use Case 11: Single File RAG Query with Citations")
    try:
        pdf_path = "Project_Alpha_Summary.pdf"
        pdf_content = ["Project Alpha focuses on a new renewable energy source.",
                       "Key Objective: Achieve 20% efficiency by Q4.", "Status: At 15%, a 5% increase from Q2.",
                       "Conclusion: On track, but needs more funding for Phase 2."]
        create_dummy_pdf(pdf_path, "Project Alpha - Q3 Report Summary", pdf_content)
        prompt = "Summarize the attached document. What is the key objective and the final conclusion?"
        print(f"\nUser Query: {prompt}\n")

        with client.upload_and_manage_files([pdf_path]) as uploaded_files:
            stream = client.stream_chat_completion(
                model=MODEL,
                messages=[{"role": "user", "content": prompt}],
                uploaded_files=uploaded_files
            )
            print("Assistant Response (streaming):\n---")
            answer = stream_and_print_response(stream)
            print("---")
            print_source_references_if_any(answer, uploaded_files)

        os.remove(pdf_path)

    except Exception as e:
        print(f"An error occurred: {e}")


def use_case_12_multi_file_conversation(client: OpenWebUIClient):
    """12. Upload multiple files and hold a conversation with citations."""
    print_header("Use Case 12: Multi-File RAG Conversational Query")
    try:
        pdf_path = "Project_Helios_Plan.pdf"
        create_dummy_pdf(pdf_path, "Project Helios - Official Plan",
                         ["Team Lead: Dr. Aris Thorne", "Budget: $500,000", "Goal: Solar-powered water purification."])
        docx_path = "Helios_Meeting_Minutes.docx"
        create_dummy_docx(docx_path, "Meeting Minutes - Helios Kick-off",
                          ["Attendees: Dr. Aris Thorne, Ben Carter.", "Decision: Ben Carter is lead engineer.",
                           "Risk: Supply chain delays."])

        with client.upload_and_manage_files([pdf_path, docx_path]) as uploaded_files:
            builder = OpenWebUIMessageBuilder()
            prompt1 = "Based on both documents, who is the team lead and who is the lead engineer?"
            print(f"User Query 1: {prompt1}\n")
            builder.add_user_message(prompt1)
            stream1 = client.stream_chat_completion(
                model=MODEL,
                messages=builder.build(),
                uploaded_files=uploaded_files)
            print("Assistant Response 1 (streaming):\n---")
            answer1 = stream_and_print_response(stream1)
            print("---")
            print_source_references_if_any(answer1, uploaded_files)
            builder.add_assistant_message(answer1)

            prompt2 = "\nWhat is the main risk identified for that project?"
            print(f"User Query 2: {prompt2}\n")
            builder.add_user_message(prompt2)
            stream2 = client.stream_chat_completion(
                model=MODEL,
                messages=builder.build(),
                uploaded_files=uploaded_files)
            print("Assistant Response 2 (streaming):\n---")
            answer2 = stream_and_print_response(stream2)
            print("---")
            print_source_references_if_any(answer2, uploaded_files)

        os.remove(pdf_path)
        os.remove(docx_path)

    except Exception as e:
        print(f"An error occurred: {e}")


def main():
    """Main function to run all tutorial use cases."""
    if not all([BASE_URL, API_KEY, MODEL]):
        print("Error: Missing required environment variables.")
        print("Please set OPENWEBUI_URL, OPENWEBUI_API_KEY, and OPENWEBUI_MODEL.")
        sys.exit(1)

    try:
        client = OpenWebUIClient(base_url=BASE_URL, api_key=API_KEY)

        # Run all use cases in a logical learning progression
        use_case_1_list_models(client)
        use_case_2_simple_query(client)
        use_case_3_streaming_query(client)
        use_case_4_conversational_context(client)
        use_case_5_system_prompt_persona(client)
        use_case_6_robust_json_extraction(client)
        use_case_7_json_mode(client)
        use_case_8_generation_parameters(client)
        use_case_9_single_image_query(client)
        use_case_10_multi_image_query(client)
        use_case_11_single_file_query(client)
        use_case_12_multi_file_conversation(client)

    except Exception as e:
        print(f"A critical error occurred: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()